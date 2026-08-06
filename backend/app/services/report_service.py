"""Generates real PDF and DOCX forensic reports for a case (Module 14).

Pulls directly from the same `Case`/`Evidence`/`AIResult` rows Module 13 already persists —
no separate report-specific data model, no re-running the AI pipeline. A report is a
formatted view of data that already exists and was already computed honestly (see
app/api/routes/cases.py's docstring on why verdicts are always server-computed).
"""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from database.models import Case

DISCLAIMER = (
    "This report was generated automatically by TRACER. Verdicts are produced by an "
    "automated detector (Module 5's offline evaluation measured 86% binary accuracy, "
    "0.95 AUC on its test set) and are not a substitute for expert human review. Treat "
    "every verdict below as an investigative lead, not a final determination."
)


def _summary_counts(case: Case) -> dict:
    total = len(case.evidence)
    clean = sum(1 for e in case.evidence if e.ai_result and e.ai_result.verdict.value == "clean")
    adversarial = sum(
        1 for e in case.evidence if e.ai_result and e.ai_result.verdict.value == "adversarial"
    )
    return {"total": total, "clean": clean, "adversarial": adversarial}


def generate_pdf_report(case: Case, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    body_style = styles["BodyText"]
    disclaimer_style = ParagraphStyle(
        "Disclaimer", parent=body_style, fontSize=8, textColor=colors.grey
    )

    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    story = [
        Paragraph("TRACER Forensic Report", title_style),
        Spacer(1, 12),
        Paragraph(f"{case.case_number} — {case.title}", heading_style),
        Paragraph(f"Status: {case.status.value.replace('_', ' ').title()}", body_style),
    ]
    if case.description:
        story.append(Paragraph(f"Description: {case.description}", body_style))
    story.append(
        Paragraph(
            f"Case created: {case.created_at.strftime('%Y-%m-%d %H:%M UTC')} · "
            f"Report generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            body_style,
        )
    )
    story.append(Spacer(1, 12))

    counts = _summary_counts(case)
    summary_table = Table(
        [
            ["Total Evidence", "Clean", "Adversarial"],
            [str(counts["total"]), str(counts["clean"]), str(counts["adversarial"])],
        ],
        colWidths=[1.8 * inch] * 3,
    )
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#232733")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.append(summary_table)
    story.append(Spacer(1, 20))

    for evidence in sorted(case.evidence, key=lambda e: e.uploaded_at):
        story.append(Paragraph(evidence.original_filename, heading_style))
        story.append(Paragraph(f"SHA-256: {evidence.sha256_hash}", disclaimer_style))
        story.append(
            Paragraph(
                f"Uploaded: {evidence.uploaded_at.strftime('%Y-%m-%d %H:%M UTC')}", body_style
            )
        )

        result = evidence.ai_result
        if result is None:
            story.append(Paragraph("No AI result recorded for this evidence.", body_style))
            story.append(Spacer(1, 16))
            continue

        story.append(
            Paragraph(
                f"<b>Verdict: {result.verdict.value.upper()}</b> "
                f"({result.confidence * 100:.1f}% confidence)",
                body_style,
            )
        )
        if result.attack_type:
            story.append(
                Paragraph(
                    f"Attack type: {result.attack_type} "
                    f"({(result.attack_type_confidence or 0) * 100:.1f}% confidence)",
                    body_style,
                )
            )
        story.append(
            Paragraph(f"Attribution method: {result.attribution_method}", body_style)
        )

        if result.attribution_heatmap_path and Path(result.attribution_heatmap_path).exists():
            story.append(Spacer(1, 6))
            story.append(RLImage(result.attribution_heatmap_path, width=3 * inch, height=3 * inch))

        story.append(Spacer(1, 6))
        story.append(Paragraph(result.explanation_summary, body_style))
        if result.explanation_details:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(d, body_style)) for d in result.explanation_details],
                    bulletType="bullet",
                )
            )
        story.append(Spacer(1, 20))

    story.append(Spacer(1, 12))
    story.append(Paragraph(DISCLAIMER, disclaimer_style))

    doc.build(story)


def generate_docx_report(case: Case, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    document.add_heading("TRACER Forensic Report", level=0)
    document.add_heading(f"{case.case_number} — {case.title}", level=1)
    document.add_paragraph(f"Status: {case.status.value.replace('_', ' ').title()}")
    if case.description:
        document.add_paragraph(f"Description: {case.description}")
    document.add_paragraph(
        f"Case created: {case.created_at.strftime('%Y-%m-%d %H:%M UTC')} · "
        f"Report generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    counts = _summary_counts(case)
    table = document.add_table(rows=2, cols=3)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Total Evidence"
    header_cells[1].text = "Clean"
    header_cells[2].text = "Adversarial"
    value_cells = table.rows[1].cells
    value_cells[0].text = str(counts["total"])
    value_cells[1].text = str(counts["clean"])
    value_cells[2].text = str(counts["adversarial"])

    for evidence in sorted(case.evidence, key=lambda e: e.uploaded_at):
        document.add_heading(evidence.original_filename, level=2)
        hash_paragraph = document.add_paragraph(f"SHA-256: {evidence.sha256_hash}")
        hash_paragraph.runs[0].font.size = Pt(8)
        document.add_paragraph(f"Uploaded: {evidence.uploaded_at.strftime('%Y-%m-%d %H:%M UTC')}")

        result = evidence.ai_result
        if result is None:
            document.add_paragraph("No AI result recorded for this evidence.")
            continue

        verdict_paragraph = document.add_paragraph()
        verdict_run = verdict_paragraph.add_run(
            f"Verdict: {result.verdict.value.upper()} ({result.confidence * 100:.1f}% confidence)"
        )
        verdict_run.bold = True

        if result.attack_type:
            document.add_paragraph(
                f"Attack type: {result.attack_type} "
                f"({(result.attack_type_confidence or 0) * 100:.1f}% confidence)"
            )
        document.add_paragraph(f"Attribution method: {result.attribution_method}")

        if result.attribution_heatmap_path and Path(result.attribution_heatmap_path).exists():
            document.add_picture(result.attribution_heatmap_path, width=Inches(3))

        document.add_paragraph(result.explanation_summary)
        for detail in result.explanation_details:
            document.add_paragraph(detail, style="List Bullet")

    document.add_paragraph()
    disclaimer_paragraph = document.add_paragraph(DISCLAIMER)
    disclaimer_paragraph.runs[0].font.size = Pt(8)

    document.save(str(output_path))
